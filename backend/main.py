from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pypdf import PdfReader
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
import docx2txt
import io
import os
import json

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise Exception("OPENAI_API_KEY not found. Check backend/.env file.")

client = OpenAI(api_key=api_key)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def extract_resume_text(file_bytes, filename):
    filename = filename.lower()

    if filename.endswith(".pdf"):
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text

    if filename.endswith(".docx"):
        temp_file = "temp_resume.docx"

        with open(temp_file, "wb") as f:
            f.write(file_bytes)

        text = docx2txt.process(temp_file)
        os.remove(temp_file)

        return text

    raise Exception("Only PDF and DOCX files are supported right now")


def call_openai_json(prompt, system_message, temperature=0.3):
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        temperature=temperature,
    )

    ai_text = response.choices[0].message.content
    return json.loads(ai_text)


@app.get("/")
def home():
    return {"message": "AI Resume Analyzer backend is running"}


@app.post("/analyze-resume")
async def analyze_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    try:
        contents = await file.read()
        resume_text = extract_resume_text(contents, file.filename)

        prompt = f"""
Analyze this resume against the job description.

Return ONLY valid JSON with these fields:
{{
  "match_score": number,
  "overall_summary": "string",
  "strong_matches": ["string"],
  "missing_skills": ["string"],
  "resume_improvements": ["string"],
  "interview_questions": ["string"]
}}

Resume:
{resume_text[:6000]}

Job Description:
{job_description[:4000]}
"""

        ai_result = call_openai_json(
            prompt,
            "You are an expert ATS resume analyst.",
            temperature=0.2,
        )

        return {
            "filename": file.filename,
            "resume_preview": resume_text[:500],
            **ai_result
        }

    except Exception as e:
        return {"error": str(e)}


@app.post("/generate-cover-letter")
async def generate_cover_letter(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    try:
        contents = await file.read()
        resume_text = extract_resume_text(contents, file.filename)

        prompt = f"""
Write a professional, concise cover letter for this job.

Rules:
- Do not invent fake experience.
- Use the resume information only.
- Make it tailored to the job description.
- Keep it under 350 words.
- Use a confident but natural tone.

Resume:
{resume_text[:6000]}

Job Description:
{job_description[:4000]}
"""

        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are an expert career coach and cover letter writer."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
        )

        return {
            "filename": file.filename,
            "cover_letter": response.choices[0].message.content
        }

    except Exception as e:
        return {"error": str(e)}


@app.post("/rewrite-resume")
async def rewrite_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    try:
        contents = await file.read()
        resume_text = extract_resume_text(contents, file.filename)

        prompt = f"""
Rewrite and optimize this resume for the job description.

Rules:
- Do not invent fake companies, fake dates, fake degrees, or fake experience.
- Keep the rewritten content based only on the resume.
- Make it ATS-friendly.
- Use strong action verbs.
- Tailor wording to the job description.
- Improve clarity, impact, and keyword alignment.

Return ONLY valid JSON with these fields:
{{
  "optimized_summary": "string",
  "optimized_skills": ["string"],
  "optimized_bullets": ["string"],
  "ats_keywords_to_add": ["string"]
}}

Resume:
{resume_text[:6000]}

Job Description:
{job_description[:4000]}
"""

        rewrite_result = call_openai_json(
            prompt,
            "You are an expert resume writer and ATS optimization specialist.",
            temperature=0.3,
        )

        return {
            "filename": file.filename,
            **rewrite_result
        }

    except Exception as e:
        return {"error": str(e)}


@app.post("/generate-tailored-resume")
async def generate_tailored_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    mode: str = Form("ATS Optimized")
):
    try:
        contents = await file.read()
        resume_text = extract_resume_text(contents, file.filename)

        prompt = f"""
Create a complete ATS-tailored resume based on the original resume and job description.

Important rules:
- Do not invent fake companies.
- Do not invent fake degrees.
- Do not invent fake dates.
- Do not claim expert-level skills unless clearly supported.
- You may add truthful related keywords naturally where the resume supports them.
- If a keyword is not strongly supported, phrase it carefully as exposure/familiarity when appropriate.
- Make the resume strong, ATS-friendly, and recruiter-readable.
- Mode: {mode}

Return ONLY valid JSON with this exact structure:
{{
  "name": "string",
  "contact": "string",
  "target_title": "string",
  "professional_summary": "string",
  "skills": ["string"],
  "experience": [
    {{
      "title": "string",
      "company": "string",
      "location": "string",
      "dates": "string",
      "bullets": ["string"]
    }}
  ],
  "projects": [
    {{
      "name": "string",
      "bullets": ["string"]
    }}
  ],
  "education": ["string"],
  "certifications": ["string"]
}}

Original Resume:
{resume_text[:8000]}

Job Description:
{job_description[:5000]}
"""

        tailored = call_openai_json(
            prompt,
            "You are an expert ATS resume writer. You create truthful, job-targeted resumes.",
            temperature=0.3,
        )

        doc = Document()

        doc.add_heading(tailored.get("name", "Tailored Resume"), level=0)

        if tailored.get("contact"):
            doc.add_paragraph(tailored["contact"])

        if tailored.get("target_title"):
            doc.add_heading(tailored["target_title"], level=1)

        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(tailored.get("professional_summary", ""))

        doc.add_heading("Technical Skills", level=1)
        skills = tailored.get("skills", [])
        if skills:
            doc.add_paragraph(", ".join(skills))

        doc.add_heading("Professional Experience", level=1)
        for exp in tailored.get("experience", []):
            title_line = f"{exp.get('title', '')} | {exp.get('company', '')}"
            doc.add_heading(title_line.strip(" | "), level=2)

            meta = " | ".join(
                item for item in [
                    exp.get("location", ""),
                    exp.get("dates", "")
                ] if item
            )

            if meta:
                doc.add_paragraph(meta)

            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        projects = tailored.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=1)
            for project in projects:
                doc.add_heading(project.get("name", "Project"), level=2)
                for bullet in project.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

        education = tailored.get("education", [])
        if education:
            doc.add_heading("Education", level=1)
            for item in education:
                doc.add_paragraph(item, style="List Bullet")

        certifications = tailored.get("certifications", [])
        if certifications:
            doc.add_heading("Certifications", level=1)
            for item in certifications:
                doc.add_paragraph(item, style="List Bullet")

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=ATS_Tailored_Resume.docx"
            }
        )

    except Exception as e:
        return {"error": str(e)}